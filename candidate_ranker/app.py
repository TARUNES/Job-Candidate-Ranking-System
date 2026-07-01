import io
import json
import time
import gzip
import csv
import numpy as np
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
import uvicorn
import docx

# Imports from candidate_ranker pipeline
from src.embeddings import get_sentence_embeddings
from src.heuristics import MismatchDetector, check_honeypots_and_filters, compute_soft_penalty
from src.jd_parser import JDParser
from src.schema_validator import validate_candidate
from src.scoring import compute_non_semantic_score
from rank import extract_candidate_text, select_reasoning

app = FastAPI(title="Candidate Discovery and Ranking System")

# Constants matching rank.py exactly
SHORTLIST_SIZE = 2000
SEMANTIC_WEIGHT = 0.45
NON_SEMANTIC_WEIGHT = 0.55

# Helper functions
def get_jd_text_from_bytes(file_bytes, filename):
    if filename.endswith('.docx'):
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    else:
        # Assume plain text
        return file_bytes.decode('utf-8', errors='ignore')

def stream_candidates_from_bytes(file_bytes, filename):
    if filename.endswith('.gz'):
        fileobj = io.BytesIO(file_bytes)
        with gzip.GzipFile(fileobj=fileobj, mode='rb') as gz:
            with io.TextIOWrapper(gz, encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        yield json.loads(line)
    elif filename.endswith('.json'):
        content = file_bytes.decode('utf-8', errors='ignore')
        data = json.loads(content)
        if isinstance(data, list):
            for item in data:
                yield item
        else:
            yield data
    else:
        # Assume plain JSONLines (.jsonl)
        content = file_bytes.decode('utf-8', errors='ignore')
        for line in content.splitlines():
            if line.strip():
                yield json.loads(line)

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        return f.read()

async def run_ranking_pipeline_generator(jd_bytes, jd_name, candidates_bytes, candidates_name):
    t_start = time.perf_counter()
    try:
        # 1. Load ML Models
        yield json.dumps({"step": 1, "status": "running", "message": "Loading SentenceTransformer (all-MiniLM-L6-v2) and Google Flan-T5-base from local cache..."}) + "\n"
        from src.embeddings import LocalEncoder
        LocalEncoder.get_model()
        from src.jd_parser import _LLMSingleton
        _LLMSingleton._load()
        yield json.dumps({"step": 1, "status": "success", "message": "Models initialized. Offline inference engines ready."}) + "\n"

        # 2. Parse JD
        yield json.dumps({"step": 2, "status": "running", "message": "Parsing Job Description..."}) + "\n"
        jd_text = get_jd_text_from_bytes(jd_bytes, jd_name)
        jd_profile = JDParser.parse(jd_text)
        yield json.dumps({
            "step": 2, 
            "status": "success", 
            "message": f"Job Description parsed: found {len(jd_profile.required_skills)} required skills, {len(jd_profile.preferred_skills)} preferred skills, ideal experience {jd_profile.ideal_years:.1f} YOE."
        }) + "\n"

        # 3. Embed JD
        yield json.dumps({"step": 3, "status": "running", "message": "Generating Job Description embedding..."}) + "\n"
        jd_vector = get_sentence_embeddings([jd_text])[0]
        yield json.dumps({"step": 3, "status": "success", "message": "JD Embedded successfully."}) + "\n"

        # 4. Stream & Filter candidates
        yield json.dumps({"step": 4, "status": "running", "message": "Streaming candidates and running hard-filters..."}) + "\n"
        valid_candidates = []
        total_count = 0
        invalid_count = 0
        filtered_count = 0

        for cand in stream_candidates_from_bytes(candidates_bytes, candidates_name):
            total_count += 1
            is_valid, reason = validate_candidate(cand)
            if not is_valid:
                invalid_count += 1
                continue
            if check_honeypots_and_filters(cand):
                filtered_count += 1
                continue
            valid_candidates.append(cand)

        yield json.dumps({
            "step": 4, 
            "status": "success", 
            "message": f"Streamed {total_count} candidates: {len(valid_candidates)} valid profiles, {invalid_count} schema invalid, {filtered_count} hard-filtered."
        }) + "\n"

        if not valid_candidates:
            yield json.dumps({"status": "error", "message": "No valid candidates passed the schema validation and hard heuristic filters."}) + "\n"
            return

        # 5. TF-IDF mismatch detection
        yield json.dumps({"step": 5, "status": "running", "message": "Running data-driven domain mismatch detector..."}) + "\n"
        detector = MismatchDetector()
        mismatched_ids = detector.detect(valid_candidates)
        mismatch_count = len(mismatched_ids)
        if mismatched_ids:
            valid_candidates = [c for c in valid_candidates if c["candidate_id"] not in mismatched_ids]
        yield json.dumps({
            "step": 5, 
            "status": "success", 
            "message": f"Coherence analysis completed: {mismatch_count} profiles flagged for domain mismatches."
        }) + "\n"

        # 6. Pre-ranking (non-semantic score)
        yield json.dumps({"step": 6, "status": "running", "message": "Screen and pre-rank remaining profiles..."}) + "\n"
        ns_scores = []
        sp_scores = []
        for cand in valid_candidates:
            ns_scores.append(float(compute_non_semantic_score(cand, jd_profile)))
            sp_scores.append(float(compute_soft_penalty(cand)))

        # Sort
        ns_penalised = [ns * sp for ns, sp in zip(ns_scores, sp_scores)]
        ranked_indices = sorted(
            range(len(valid_candidates)),
            key=lambda i: ns_penalised[i],
            reverse=True,
        )
        actual_shortlist_size = min(SHORTLIST_SIZE, len(valid_candidates))
        shortlist_indices = ranked_indices[:actual_shortlist_size]

        shortlist_candidates = [valid_candidates[i] for i in shortlist_indices]
        shortlist_ns = np.array([ns_scores[i] for i in shortlist_indices])
        shortlist_sp = np.array([sp_scores[i] for i in shortlist_indices])
        yield json.dumps({
            "step": 6, 
            "status": "success", 
            "message": f"Pre-ranked candidates. Shortlisted top {len(shortlist_candidates)} for deep embedding."
        }) + "\n"

        # 7. Embed shortlist
        yield json.dumps({"step": 7, "status": "running", "message": "Extracting texts and encoding semantic vectors..."}) + "\n"
        candidate_texts = [extract_candidate_text(c) for c in shortlist_candidates]
        cand_vectors = get_sentence_embeddings(candidate_texts)
        yield json.dumps({"step": 7, "status": "success", "message": "Candidates encoded successfully."}) + "\n"

        # 8. Cosine similarity and hybrid calculation
        yield json.dumps({"step": 8, "status": "running", "message": "Performing hybrid scoring..."}) + "\n"
        jd_norm = np.linalg.norm(jd_vector)
        cand_norms = np.linalg.norm(cand_vectors, axis=1)
        safe_denom = jd_norm * cand_norms
        safe_denom[safe_denom == 0] = 1.0
        sem_scores = np.dot(cand_vectors, jd_vector) / safe_denom

        ns_penalised_arr = shortlist_ns * shortlist_sp
        total_scores = SEMANTIC_WEIGHT * sem_scores + NON_SEMANTIC_WEIGHT * ns_penalised_arr
        yield json.dumps({"step": 8, "status": "success", "message": "Hybrid scores computed."}) + "\n"

        # 9. Reasoning & Results
        yield json.dumps({"step": 9, "status": "running", "message": "Generating grounded reasoning justifications..."}) + "\n"
        results = []
        for i, cand in enumerate(shortlist_candidates):
            reasoning = select_reasoning(
                cand, float(sem_scores[i]), float(total_scores[i]), jd_profile
            )
            results.append({
                "candidate_id": cand["candidate_id"],
                "score": round(float(total_scores[i]), 4),
                "reasoning": reasoning,
                "profile_details": cand  # Store full profile for verification
            })

        # Sort top list (by score descending, candidate_id ascending)
        results.sort(key=lambda x: (-x["score"], x["candidate_id"]))
        yield json.dumps({"step": 9, "status": "success", "message": "Sorting completed."}) + "\n"

        elapsed_time = time.perf_counter() - t_start
        
        # Format the top-100 results for rendering
        table_results = []
        for rank_val, res in enumerate(results, start=1):
            table_results.append({
                "rank": rank_val,
                "candidate_id": res["candidate_id"],
                "score": res["score"],
                "reasoning": res["reasoning"]
            })

        # Build CSV data
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(["candidate_id", "rank", "score", "reasoning"])
        for rank_val, res in enumerate(results[:100], start=1):
            writer.writerow([
                res["candidate_id"],
                rank_val,
                res["score"],
                res["reasoning"]
            ])
        csv_data = csv_buffer.getvalue()

        # Send final completion message
        yield json.dumps({
            "status": "completed",
            "elapsed": round(elapsed_time, 2),
            "metrics": {
                "total_input": total_count,
                "invalid": invalid_count,
                "filtered": filtered_count,
                "mismatched": mismatch_count,
                "ranked": len(results)
            },
            "results": table_results,
            "raw_profiles": {res["candidate_id"]: res["profile_details"] for res in results},
            "csv_data": csv_data
        }) + "\n"
        
    except Exception as e:
        yield json.dumps({"status": "error", "message": str(e)}) + "\n"

@app.post("/rank")
async def run_ranking(
    jd_file: UploadFile = File(...),
    candidates_file: UploadFile = File(...)
):
    try:
        jd_bytes = await jd_file.read()
        candidates_bytes = await candidates_file.read()
        
        return StreamingResponse(
            run_ranking_pipeline_generator(jd_bytes, jd_file.filename, candidates_bytes, candidates_file.filename),
            media_type="text/event-stream"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8501, reload=True)
