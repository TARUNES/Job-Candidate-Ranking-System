import os
import json
import gzip
import docx

def load_docx_text(file_path):
    """
    Parses a Microsoft Word Document (.docx) file and returns its textual content.
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def stream_candidates(file_path):
    """
    Generator function that streams candidate profiles line-by-line from a gzipped or plain text JSONLines file to optimize memory consumption.
    """
    if file_path.endswith('.gz'):
        with gzip.open(file_path, 'rt', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    yield json.loads(line)
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    yield json.loads(line)
